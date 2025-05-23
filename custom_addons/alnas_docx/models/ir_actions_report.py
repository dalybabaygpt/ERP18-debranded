import base64
import zipfile
import os
import subprocess
import tempfile
import shutil
from io import BytesIO
from functools import partial
from docxtpl import DocxTemplate, InlineImage
from docxcompose.composer import Composer
from docx import Document
from docx.shared import Mm
from bs4 import BeautifulSoup
from num2words import num2words
from babel.dates import format_date
from babel.numbers import format_currency
from htmldocx import HtmlToDocx

from odoo import _, api, fields, models
from odoo.tools.safe_eval import safe_eval, time
from odoo.exceptions import ValidationError


class IrActionsReport(models.Model):
    _inherit = "ir.actions.report"

    report_type = fields.Selection(
        selection_add=[("docx", "DOCX")], ondelete={"docx": "cascade"}
    ) # add docx type
    report_docx_template = fields.Binary(string="Report DOCX Template")
    report_docx_template_name = fields.Char(string="Report DOCX Template Name")
    docx_merge_mode = fields.Selection(
        [("composer", "Composer"), ("zip", "Zip"), ("pdf", "PDF")],
        string="DOCX Mode",
        default="composer",
    )

    @api.constrains("report_type")
    def _check_report_type(self):
        for rec in self:
            if (
                rec.report_type == "docx"
                and not rec.report_docx_template
                and not rec.report_docx_template_name.endswith(".docx")
            ):
                raise ValidationError(_("Please upload a DOCX template."))

    def _render_docx(self, report_ref, docids, data):
        report = self._get_report(report_ref)
        template = report.report_docx_template

        if not template:
            raise ValueError("No DOCX template found.")

        doc_template = DocxTemplate(BytesIO(base64.b64decode(template)))
        doc_obj = self.env[report.model].browse(docids)
        render_image = partial(self._render_image, doc_template)
        html2docx = partial(self._render_html_as_subdoc, doc_template)

        context = {
            "spelled_out": self._spelled_out,
            "parsehtml": self._parse_html,
            "formatdate": self._formatdate,
            "company": self.env.company,
            "lang": self._context.get("lang", "id_ID"),
            "sysdate": fields.Datetime.now(),
            "render_image": render_image,
            "html2docx": html2docx,
            "convert_currency": self._convert_currency,
            "formatabs": self._format_abs
        }

        if report.docx_merge_mode == "composer":
            return self._render_composer_mode(doc_template, doc_obj, data, context)
        elif report.docx_merge_mode == "zip":
            return self._render_zip_mode(
                doc_template,
                doc_obj,
                data,
                context,
                report_name=report.print_report_name,
            )
        else:
            return self._render_docx_to_pdf_mode(doc_template, doc_obj, data, context)

    def _render_composer_mode(self, doc_template, doc_obj, data, context):
        master_doc = Document()
        composer = Composer(master_doc)

        for idx, obj in enumerate(doc_obj):
            context = {
                **context,
                "docs": obj,
                "data": data
            }

            temp = BytesIO()
            doc_template.render(context)
            doc_template.save(temp)
            temp.seek(0)

            if len(doc_obj) == 1:
                return temp.read(), 'docx'
            else:
                if idx == 0:
                    master_doc = Document(temp)
                    composer = Composer(master_doc)
                else:
                    doc_to_append = Document(temp)
                    master_doc.add_page_break()
                    composer.append(doc_to_append)

        temp_output = BytesIO()
        composer.save(temp_output)
        temp_output.seek(0)

        return temp_output.read(), 'docx'

    def _render_zip_mode(
        self, doc_template, doc_obj, data, context, report_name="report"
    ):
        docx_files = []

        for obj in doc_obj:
            context = {
                **context,
                "docs": obj,
                "data": data
            }

            temp = BytesIO()
            doc_template.render(context)
            doc_template.save(temp)
            temp.seek(0)
            docx_files.append(temp.read())

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zip_file:
            for idx, docx_file in enumerate(docx_files):
                name = safe_eval(report_name, {"object": doc_obj[idx], "time": time})
                filename = "%s.%s" % (name, "docx")
                zip_file.writestr(filename, docx_file)

        zip_buffer.seek(0)

        return zip_buffer.read(), 'zip'

    def _render_docx_to_pdf_mode(self, doc_template, doc_obj, data, context):
        docx_file, _ = self._render_composer_mode(doc_template, doc_obj, data, context)
        temp_dir = tempfile.mkdtemp()
        os.makedirs(temp_dir, exist_ok=True)

        try:
            docx_file_path = os.path.join(temp_dir, 'document.docx')
            with open(docx_file_path, 'wb') as f:
                f.write(docx_file)

            pdf_file_path = self.convert_file_to_pdf(docx_file_path, temp_dir)

            if not pdf_file_path:
                raise Exception('PDF conversion failed.')

            with open(pdf_file_path, 'rb') as pdf_file:
                pdf_bytes = BytesIO(pdf_file.read())

        finally:
            shutil.rmtree(temp_dir)

        return pdf_bytes.read(), 'pdf'

    def convert_file_to_pdf(self, file_path, output_dir):
        librepath = self._get_libreoffice_path()
        subprocess.run([librepath, '--headless', '--convert-to', 'pdf', '--outdir', output_dir, file_path])
        pdf_file_name = os.path.splitext(os.path.basename(file_path))[0] + '.pdf' 
        pdf_file_path = os.path.join(output_dir, pdf_file_name)        
        if os.path.exists(pdf_file_path):
            return pdf_file_path
        else:
            return None

    def _get_libreoffice_path(self):
        libreoffice = self.env.ref('alnas_docx.default_libreoffice_path')
        if not libreoffice and not libreoffice.value:
            raise ValidationError('Libreoffice path doesnt exits, \n \
                please set in Settings => Technical => Parameters => System Parameters => default_libreoffice_path')
            
        return libreoffice.value

    @staticmethod
    def _render_image(tpl, imgb64, width=None, height=None):
        width = Mm(width) if width else None
        height = Mm(height) if height else None
        
        if not imgb64:
            return ''

        image_stream = BytesIO(base64.b64decode(imgb64))
        return InlineImage(
            tpl, image_descriptor=image_stream, width=width, height=height
        )

    @staticmethod
    def _parse_html(html):
        if not html:
            return ""
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text()

    @staticmethod
    def _formatdate(date_required=fields.Datetime.today(), format="full", lang="id_ID", **kwargs):
        return format_date(date=date_required, format=format, locale=lang, **kwargs)

    @staticmethod
    def _spelled_out(number, lang="id_ID", **kwargs):
        return num2words(number=number, lang=lang, **kwargs)
    
    @staticmethod
    def _convert_currency(number, currency_field, locale='id_ID', **kwargs):
        return format_currency(number=number, currency=currency_field.name, locale=locale, **kwargs)

    @staticmethod
    def _format_abs(number):
        return abs(number)

    @staticmethod
    def _render_html_as_subdoc(tpl, html_code=None):
        if not (
            isinstance(html_code, str)
            and bool(BeautifulSoup(html_code, "html.parser").find())
        ):
            return ""

        temp = BytesIO()
        desc_document = Document()
        new_parser = HtmlToDocx()
        new_parser.add_html_to_document(html_code, desc_document)
        desc_document.save(temp)
        temp.seek(0)
        return tpl.new_subdoc(temp)
